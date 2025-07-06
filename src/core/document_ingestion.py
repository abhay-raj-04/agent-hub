import os
from ..types.state import ResumeCoPilotState

def ingest_documents_node(state: ResumeCoPilotState, file_path_resume=None, file_path_jd=None):
    """
    Extract text from uploaded files (resume, JD). Update state with extracted text.
    """
    from PyPDF2 import PdfReader
    from docx import Document

    try:
        # In a real app, file paths would be provided. Here, we check if text is already present.
        if not state.get("resume_text") and file_path_resume:
            ext = os.path.splitext(file_path_resume)[1].lower()
            if ext == ".pdf":
                with open(file_path_resume, "rb") as f:
                    reader = PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() or ""
                    state["resume_text"] = text
            elif ext == ".docx":
                doc = Document(file_path_resume)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                state["resume_text"] = text
            else:
                state["error_message"] = "Unsupported resume file type."
        if not state.get("job_description_text") and file_path_jd:
            ext = os.path.splitext(file_path_jd)[1].lower()
            if ext == ".pdf":
                with open(file_path_jd, "rb") as f:
                    reader = PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() or ""
                    state["job_description_text"] = text
            elif ext == ".docx":
                doc = Document(file_path_jd)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                state["job_description_text"] = text
            elif ext == ".txt":
                with open(file_path_jd, "r", encoding="utf-8") as f:
                    state["job_description_text"] = f.read()
            else:
                state["error_message"] = "Unsupported job description file type."
    except Exception as e:
        state["error_message"] = f"Error processing documents: {e}"
    return state 