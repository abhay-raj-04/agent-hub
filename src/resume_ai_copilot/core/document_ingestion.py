import os
from ..types.state import ResumeCoPilotState

def ingest_documents_node(state: ResumeCoPilotState, file_path_resume=None, file_path_jd=None):
    """Extracts text from uploaded files (PDF, DOCX, TXT, images)."""
    from PyPDF2 import PdfReader
    from docx import Document
    from PIL import Image
    import pytesseract

    def extract_text_from_image(image_path):
        try:
            img = Image.open(image_path)
            text = pytesseract.image_to_string(img)
            return text
        except Exception as e:
            return f"[Error extracting text from image: {e}]"

    try:
        # In a real app, file paths would be provided. Here, we check if text is already present.
        if not state.get("resume_text") and file_path_resume:
            ext = os.path.splitext(file_path_resume)[1].lower()
            if ext == ".pdf":
                with open(file_path_resume, "rb") as f:
                    reader = PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() or ""
                    state["resume_text"] = text
            elif ext == ".docx":
                doc = Document(file_path_resume)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                state["resume_text"] = text
            elif ext in [".png", ".jpg", ".jpeg"]:
                state["resume_text"] = extract_text_from_image(file_path_resume)
            else:
                state["error_message"] = "Unsupported resume file type."
        if not state.get("job_description_text") and file_path_jd:
            ext = os.path.splitext(file_path_jd)[1].lower()
            if ext == ".pdf":
                with open(file_path_jd, "rb") as f:
                    reader = PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() or ""
                    state["job_description_text"] = text
            elif ext == ".docx":
                doc = Document(file_path_jd)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                state["job_description_text"] = text
            elif ext == ".txt":
                with open(file_path_jd, "r", encoding="utf-8") as f:
                    state["job_description_text"] = f.read()
            elif ext in [".png", ".jpg", ".jpeg"]:
                state["job_description_text"] = extract_text_from_image(file_path_jd)
            else:
                state["error_message"] = "Unsupported job description file type."
    except Exception as e:
        state["error_message"] = f"Error processing documents: {e}"
    return state 